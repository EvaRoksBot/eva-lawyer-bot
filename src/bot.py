import os
import tempfile
import logging
import shutil
import subprocess
from dotenv import load_dotenv

from docx import Document
from pdf2image import convert_from_bytes
import pytesseract

from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import (
    Application, MessageHandler, CallbackQueryHandler, CommandHandler,
    ContextTypes, filters,
)

from openai import OpenAI

load_dotenv()
log_level = os.getenv("LOG_LEVEL", "INFO").upper()
logging.basicConfig(level=getattr(logging, log_level, logging.INFO))
log = logging.getLogger("eva-lawyer-bot")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("AI_API_KEY", "")
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN") or os.getenv("TELEGRAM", "")
AI_BASE_URL = os.getenv("OPENAI_BASE_URL") or os.getenv("AI_BASE_URL")
AI_MODEL = os.getenv("AI_MODEL", "gpt-4o-mini")

if not TELEGRAM_BOT_TOKEN:
    raise RuntimeError("Не найден TELEGRAM_BOT_TOKEN или TELEGRAM (токен бота).")
if not OPENAI_API_KEY:
    raise RuntimeError("Не найден OPENAI_API_KEY или AI_API_KEY (ключ ИИ).")

client = OpenAI(api_key=OPENAI_API_KEY, base_url=AI_BASE_URL) if AI_BASE_URL else OpenAI(api_key=OPENAI_API_KEY)

SYSTEM_PROMPT = (
    "Ты юридический ассистент. Отвечай чётко и структурированно. "
    "Если не хватает данных, укажи, что требуется уточнить. Язык ответа: русский."
)


def _llm(prompt: str) -> str:
    resp = client.chat.completions.create(
        model=AI_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=800,
    )
    return resp.choices[0].message.content or ""


def _build_task_prompt(task: str, text: str) -> str:
    if task == "analyze_contract":
        return (
            "Проанализируй договор. Выдели предмет, обязанности, ответственность, "
            "штрафы, риски, способы защиты.\n\nТекст:\n" + text
        )
    if task == "risk_table":
        return (
            "Сделай таблицу рисков: Риск | Вероятность | Влияние | Митигирующие меры."\
            "\n\nТекст:\n" + text
        )
    if task == "verify_dd":
        return (
            "Сформируй чек-лист для проверки контрагента (Due Diligence) на основе текста."\
            "\n\nТекст:\n" + text
        )
    if task == "court_search":
        return (
            "Сформулируй поисковые запросы и возможные нормы/практику, связанные с текстом."\
            "\n\nТекст:\n" + text
        )
    if task == "analyze_claim":
        return (
            "Подготовь черновик ответа на претензию с опорой на текст."\
            "\n\nТекст:\n" + text
        )
    return "Суммаризируй ключевые положения:\n\n" + text


async def ocr_pdf_bytes(file_bytes: bytes) -> str:
    try:
        images = convert_from_bytes(file_bytes)
    except Exception as e:
        return f"[OCR] Не удалось конвертировать PDF (poppler?): {e}"
    parts = []
    for i, image in enumerate(images, 1):
        try:
            txt = pytesseract.image_to_string(image, lang="rus+eng")
        except pytesseract.TesseractNotFoundError:
            return "[OCR] Не найден бинарник tesseract-ocr в контейнере."
        except Exception as e:
            txt = f"[OCR] Ошибка на стр. {i}: {e}"
        parts.append(f"\nСтраница {i}:\n{txt}")
    return "\n".join(parts).strip()


def read_docx_bytes(file_bytes: bytes) -> str:
    """Extract text content from DOCX bytes.

    The bytes are written to a temporary file for processing via
    ``python-docx``. The temporary file is removed afterwards to avoid
    leaving artifacts on disk.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
        chunks = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                chunks.append(" | ".join(cells))
        return "\n".join(filter(None, chunks)).strip()
    finally:
        os.unlink(tmp_path)


def read_txt_bytes(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="replace").strip()
    except Exception:
        return file_bytes.decode("cp1251", errors="replace").strip()


def cut_telegram(text: str, limit: int = 4000) -> str:
    if len(text) <= limit:
        return text
    return text[: limit - 50] + "\n…(обрезано)"


async def cmd_diag(update: Update, _):
    bins = {
        "tesseract": shutil.which("tesseract"),
        "pdftoppm": shutil.which("pdftoppm"),
    }
    env_ok = bool(OPENAI_API_KEY and TELEGRAM_BOT_TOKEN)
    try:
        tver = subprocess.check_output(["tesseract", "--version"], text=True).splitlines()[0]
    except Exception as e:
        tver = f"err: {e}"
    msg = (
        f"✅ ENV: {env_ok}\n"
        f"🔑 AI key: {bool(OPENAI_API_KEY)} | 🤖 TG token: {bool(TELEGRAM_BOT_TOKEN)}\n"
        f"🧠 Model: {AI_MODEL}\n"
        f"📦 Binaries: {bins}\n"
        f"🖹 tesseract: {tver}\n"
    )
    await update.message.reply_text(cut_telegram(msg))


MAX_INPUT_CHARS = 18000


async def handle_any_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.document:
        return
    doc = update.message.document
    name = (doc.file_name or "").lower()
    mime = (doc.mime_type or "").lower()

    tg_file = await doc.get_file()
    file_bytes = await tg_file.download_as_bytearray()

    if "pdf" in mime or name.endswith(".pdf"):
        await update.message.reply_text("Распознаю PDF…")
        context.user_data["extracted_text"] = await ocr_pdf_bytes(bytes(file_bytes))
    elif "wordprocessingml" in mime or name.endswith(".docx"):
        await update.message.reply_text("Читаю .docx…")
        context.user_data["extracted_text"] = read_docx_bytes(bytes(file_bytes))
    elif mime.startswith("text/") or name.endswith(".txt"):
        await update.message.reply_text("Читаю .txt…")
        context.user_data["extracted_text"] = read_txt_bytes(bytes(file_bytes))
    else:
        await update.message.reply_text("Формат не поддержан. Пришлите PDF, DOCX или TXT.")
        return

    msg = f"📄 Текст получен. Символов: {len(context.user_data.get('extracted_text',''))}.\nВыберите действие:"
    kb = InlineKeyboardMarkup([
        [InlineKeyboardButton("🔍 Анализ договора", callback_data='analyze_contract')],
        [InlineKeyboardButton("⚠️ Таблица рисков", callback_data='risk_table')],
        [InlineKeyboardButton("🧾 Проверка контрагента (DD)", callback_data='verify_dd')],
        [InlineKeyboardButton("🔎 Поиск практики", callback_data='court_search')],
        [InlineKeyboardButton("✉️ Ответ на претензию", callback_data='analyze_claim')],
    ])
    await update.message.reply_text(msg, reply_markup=kb)


async def on_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    task = query.data
    text = context.user_data.get("extracted_text") or ""
    if not text:
        await query.edit_message_text("Нет текста. Пришлите PDF/DOCX/TXT.")
        return
    await query.edit_message_text("Обрабатываю запрос ИИ…")
    snippet = text if len(text) <= MAX_INPUT_CHARS else text[:MAX_INPUT_CHARS]
    prompt = _build_task_prompt(task, snippet)
    try:
        answer = _llm(prompt)
    except Exception as e:
        log.exception("LLM error")
        answer = f"Ошибка вызова ИИ: {e}"
    await query.edit_message_text(cut_telegram(answer or "Пустой ответ."))


def build_app() -> Application:
    if not TELEGRAM_BOT_TOKEN:
        raise RuntimeError("TELEGRAM_BOT_TOKEN не задан")
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.add_handler(CommandHandler("diag", cmd_diag))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_any_document))
    app.add_handler(CallbackQueryHandler(on_button))
    return app


def main():
    app = build_app()
    log.info("Bot started (long polling)…")
    app.run_polling()


if __name__ == "__main__":
    main()
